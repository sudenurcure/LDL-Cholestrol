import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

WORD_FORMAT = "Word (.docx)"

class WordPreparation:
    @staticmethod
    def generate_word(groupname):
        directory = f"all_files/{groupname}"
        all_pngs = os.listdir(directory)
        doc_name = f"{groupname}.docx"

        doc_path = os.path.join(directory, doc_name)

        # Create a Word document
        doc = Document()

        for png in all_pngs:
            testname = png.split(".")[0]
            png = directory+ "/" + png
            paragraph = doc.add_paragraph(testname+"\n")
            run = paragraph.runs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run.font.size = Pt(11)
            run.font.name = "Times New Roman"
            run.add_picture(png, width = Inches(3.92), height = Inches(0.91))
            os.remove(png)

        # Save the Word document
        doc.save(doc_path)

