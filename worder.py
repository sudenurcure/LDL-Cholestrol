import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from save_plot import insert_plot as IP

WORD_FORMAT = "Word (.docx)"

class WordPreparation:
    @staticmethod
    def generate_word(groupname):
        groupname = groupname.split(".")[0]

        directory = f"all_files/{groupname}"
        all_pngs = os.listdir(directory)
        doc_name = f"{groupname}.docx"

        doc_path = os.path.join(directory, doc_name)

        # Create a Word document
        doc = Document()

        """sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1.4)
            section.bottom_margin = Inches(0)
            section.left_margin = Inches(0.86)
            section.right_margin = Inches(0.86)
"""
        for png in all_pngs:
            testname = png.split(".")[0].replace("_"," ").replace("%", "/")
            png = directory+ "/" + png
            paragraph = doc.add_paragraph(testname+"\n")
            run = paragraph.runs[0]
            run.font.size = Pt(11)
            run.font.name = "Times New Roman"
            run.add_picture(png, width = Inches(3.92), height = Inches(0.91))
            os.remove(png)

        
        # Save the Word document
        doc.save(doc_path)

